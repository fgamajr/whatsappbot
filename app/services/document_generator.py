import os
import tempfile
from datetime import datetime
from typing import Tuple
from docx import Document
from docx.shared import Inches
import logging

logger = logging.getLogger(__name__)


class DocumentGenerator:
    def create_documents(
        self, 
        transcript: str, 
        analysis: str, 
        identifier: str
    ) -> Tuple[str, str]:
        """Create transcript and analysis documents"""
        transcript_path = self._create_transcript_document(transcript, identifier)
        analysis_path = self._create_analysis_document(analysis, identifier)
        
        return transcript_path, analysis_path
    
    def _create_transcript_document(self, transcript: str, identifier: str) -> str:
        """Create Word document with transcript"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Transcrição da Entrevista', 0)
            title.alignment = 1
            
            # Metadata
            doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d de %B de %Y às %H:%M')}")
            doc.add_paragraph(f"ID: {identifier}")
            doc.add_paragraph("")
            
            # Content
            doc.add_paragraph("Transcrição completa com timestamps:")
            doc.add_paragraph("")
            
            # Add transcript with formatting
            for line in transcript.split('\n'):
                if line.strip():
                    if line.startswith(('ENTREVISTADOR:', 'CANDIDATO:', 'LOCUTOR')):
                        para = doc.add_paragraph()
                        speaker_end = line.find(':')
                        if speaker_end > 0:
                            run = para.add_run(line[:speaker_end + 1])
                            run.bold = True
                            para.add_run(line[speaker_end + 1:])
                        else:
                            para.add_run(line)
                    else:
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph("")
            
            # Save
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"transcricao_{identifier}_{timestamp}.docx"
            doc_path = os.path.join(tempfile.gettempdir(), filename)
            doc.save(doc_path)
            
            logger.info("Transcript document created", extra={
                "file_path": doc_path
            })
            
            return doc_path
            
        except Exception as e:
            logger.error("Failed to create transcript document", extra={
                "error": str(e)
            })
            raise
    
    def _create_analysis_document(self, analysis: str, identifier: str) -> str:
        """Create Word document with analysis"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Análise da Entrevista', 0)
            title.alignment = 1
            
            # Metadata
            doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d de %B de %Y às %H:%M')}")
            doc.add_paragraph(f"ID: {identifier}")
            doc.add_paragraph("")
            
            # Content
            doc.add_paragraph("Análise estruturada baseada na transcrição:")
            doc.add_paragraph("")
            
            # Parse and format analysis
            sections = analysis.split('**')
            current_text = ""
            
            for i, section in enumerate(sections):
                if i % 2 == 1:  # Heading
                    if current_text.strip():
                        doc.add_paragraph(current_text.strip())
                        current_text = ""
                    doc.add_heading(section.strip(), level=1)
                else:
                    current_text += section
            
            if current_text.strip():
                doc.add_paragraph(current_text.strip())
            
            # Save
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"analise_{identifier}_{timestamp}.docx"
            doc_path = os.path.join(tempfile.gettempdir(), filename)
            doc.save(doc_path)
            
            logger.info("Analysis document created", extra={
                "file_path": doc_path
            })
            
            return doc_path
            
        except Exception as e:
            logger.error("Failed to create analysis document", extra={
                "error": str(e)
            })
            raise
