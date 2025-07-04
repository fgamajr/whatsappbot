import os
import json
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, List, Optional, Union
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches
import logging
from app.utils.secure_logging import SecureLogger

logger = logging.getLogger(__name__)

class ExportManager:
    """Advanced export manager supporting multiple formats and batch operations"""
    
    SUPPORTED_FORMATS = ['docx', 'pdf', 'txt', 'json', 'csv', 'xlsx', 'zip']
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def export_interview(
        self, 
        interview_data: Dict, 
        format_type: str = 'docx',
        include_analysis: bool = True,
        include_metadata: bool = True
    ) -> str:
        """Export single interview in specified format"""
        
        if format_type not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {format_type}. Supported: {self.SUPPORTED_FORMATS}")
        
        SecureLogger.safe_log_info(logger, "Starting interview export", {
            'interview_id': interview_data.get('id'),
            'format': format_type,
            'include_analysis': include_analysis
        })
        
        try:
            if format_type == 'docx':
                return self._export_to_docx(interview_data, include_analysis, include_metadata)
            elif format_type == 'pdf':
                return self._export_to_pdf(interview_data, include_analysis, include_metadata)
            elif format_type == 'txt':
                return self._export_to_txt(interview_data, include_analysis, include_metadata)
            elif format_type == 'json':
                return self._export_to_json(interview_data, include_analysis, include_metadata)
            elif format_type == 'csv':
                return self._export_to_csv(interview_data)
            elif format_type == 'xlsx':
                return self._export_to_xlsx(interview_data)
            else:
                raise ValueError(f"Export method not implemented for format: {format_type}")
                
        except Exception as e:
            SecureLogger.safe_log_error(logger, "Export failed", e, {
                'interview_id': interview_data.get('id'),
                'format': format_type
            })
            raise
    
    def export_batch(
        self, 
        interviews: List[Dict], 
        format_type: str = 'zip',
        individual_format: str = 'docx'
    ) -> str:
        """Export multiple interviews as a batch"""
        
        SecureLogger.safe_log_info(logger, "Starting batch export", {
            'interview_count': len(interviews),
            'format': format_type,
            'individual_format': individual_format
        })
        
        if format_type == 'zip':
            return self._export_batch_as_zip(interviews, individual_format)
        elif format_type == 'xlsx':
            return self._export_batch_as_xlsx(interviews)
        elif format_type == 'csv':
            return self._export_batch_as_csv(interviews)
        else:
            # Export as individual files in a directory
            return self._export_batch_as_directory(interviews, format_type)
    
    def _export_to_docx(self, interview_data: Dict, include_analysis: bool, include_metadata: bool) -> str:
        """Export to Word document with enhanced formatting"""
        doc = Document()
        
        # Document title
        title = doc.add_heading('Relatório Completo da Entrevista', 0)
        title.alignment = 1
        
        # Metadata section
        if include_metadata:
            self._add_metadata_section(doc, interview_data)
        
        # Transcript section
        self._add_transcript_section(doc, interview_data)
        
        # Analysis section
        if include_analysis and interview_data.get('analysis'):
            self._add_analysis_section(doc, interview_data)
        
        # Performance metrics
        if interview_data.get('metrics'):
            self._add_metrics_section(doc, interview_data)
        
        # Save document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevista_completa_{interview_data.get('id', 'unknown')}_{timestamp}.docx"
        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)
        
        return file_path
    
    def _export_to_pdf(self, interview_data: Dict, include_analysis: bool, include_metadata: bool) -> str:
        """Export to PDF format"""
        try:
            from docx2pdf import convert
            
            # First create DOCX
            docx_path = self._export_to_docx(interview_data, include_analysis, include_metadata)
            
            # Convert to PDF
            pdf_path = docx_path.replace('.docx', '.pdf')
            convert(docx_path, pdf_path)
            
            # Clean up temporary DOCX
            os.remove(docx_path)
            
            return pdf_path
            
        except ImportError:
            SecureLogger.safe_log_warning(logger, "PDF export requires docx2pdf package")
            raise ValueError("PDF export not available - install docx2pdf package")
    
    def _export_to_txt(self, interview_data: Dict, include_analysis: bool, include_metadata: bool) -> str:
        """Export to plain text format"""
        content = []
        
        # Header
        content.append("=" * 60)
        content.append("RELATÓRIO DA ENTREVISTA")
        content.append("=" * 60)
        content.append("")
        
        # Metadata
        if include_metadata:
            content.append("INFORMAÇÕES GERAIS:")
            content.append("-" * 20)
            content.append(f"ID da Entrevista: {interview_data.get('id', 'N/A')}")
            content.append(f"Data/Hora: {interview_data.get('created_at', 'N/A')}")
            content.append(f"Duração: {interview_data.get('duration', 'N/A')}")
            content.append(f"Status: {interview_data.get('status', 'N/A')}")
            if interview_data.get('phone_number'):
                masked_phone = SecureLogger.mask_phone_number(interview_data['phone_number'])
                content.append(f"Telefone: {masked_phone}")
            content.append("")
        
        # Transcript
        content.append("TRANSCRIÇÃO:")
        content.append("-" * 20)
        content.append(interview_data.get('transcript', 'Transcrição não disponível'))
        content.append("")
        
        # Analysis
        if include_analysis and interview_data.get('analysis'):
            content.append("ANÁLISE:")
            content.append("-" * 20)
            content.append(interview_data['analysis'])
            content.append("")
        
        # Performance metrics
        if interview_data.get('metrics'):
            content.append("MÉTRICAS DE PERFORMANCE:")
            content.append("-" * 20)
            metrics = interview_data['metrics']
            for key, value in metrics.items():
                content.append(f"{key}: {value}")
            content.append("")
        
        # Footer
        content.append("=" * 60)
        content.append(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}")
        content.append("=" * 60)
        
        # Save to file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevista_{interview_data.get('id', 'unknown')}_{timestamp}.txt"
        file_path = os.path.join(self.temp_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        return file_path
    
    def _export_to_json(self, interview_data: Dict, include_analysis: bool, include_metadata: bool) -> str:
        """Export to JSON format with structured data"""
        export_data = {
            'export_info': {
                'format': 'json',
                'exported_at': datetime.now().isoformat(),
                'version': '1.0'
            }
        }
        
        if include_metadata:
            export_data['metadata'] = {
                'interview_id': interview_data.get('id'),
                'created_at': interview_data.get('created_at'),
                'duration': interview_data.get('duration'),
                'status': interview_data.get('status'),
                'phone_number': SecureLogger.mask_phone_number(interview_data.get('phone_number', ''))
            }
        
        export_data['transcript'] = {
            'content': interview_data.get('transcript', ''),
            'chunks_total': interview_data.get('chunks_total', 0),
            'chunks_processed': interview_data.get('chunks_processed', 0)
        }
        
        if include_analysis and interview_data.get('analysis'):
            export_data['analysis'] = {
                'content': interview_data['analysis'],
                'generated_at': interview_data.get('analysis_generated_at')
            }
        
        if interview_data.get('metrics'):
            export_data['metrics'] = interview_data['metrics']
        
        # Save to file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevista_{interview_data.get('id', 'unknown')}_{timestamp}.json"
        file_path = os.path.join(self.temp_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)
        
        return file_path
    
    def _export_to_csv(self, interview_data: Dict) -> str:
        """Export to CSV format for data analysis"""
        data = {
            'interview_id': [interview_data.get('id', '')],
            'created_at': [interview_data.get('created_at', '')],
            'status': [interview_data.get('status', '')],
            'duration_minutes': [interview_data.get('duration', 0)],
            'chunks_total': [interview_data.get('chunks_total', 0)],
            'chunks_processed': [interview_data.get('chunks_processed', 0)],
            'transcript_length': [len(interview_data.get('transcript', ''))],
            'has_analysis': [bool(interview_data.get('analysis'))],
            'phone_number': [SecureLogger.mask_phone_number(interview_data.get('phone_number', ''))]
        }
        
        # Add metrics if available
        if interview_data.get('metrics'):
            for key, value in interview_data['metrics'].items():
                data[f'metric_{key}'] = [value]
        
        df = pd.DataFrame(data)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevista_data_{interview_data.get('id', 'unknown')}_{timestamp}.csv"
        file_path = os.path.join(self.temp_dir, filename)
        
        df.to_csv(file_path, index=False, encoding='utf-8')
        
        return file_path
    
    def _export_to_xlsx(self, interview_data: Dict) -> str:
        """Export to Excel format with multiple sheets"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevista_completa_{interview_data.get('id', 'unknown')}_{timestamp}.xlsx"
        file_path = os.path.join(self.temp_dir, filename)
        
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # Summary sheet
            summary_data = {
                'Campo': ['ID da Entrevista', 'Data/Hora', 'Status', 'Duração (min)', 
                         'Total de Chunks', 'Chunks Processados', 'Possui Análise'],
                'Valor': [
                    interview_data.get('id', ''),
                    interview_data.get('created_at', ''),
                    interview_data.get('status', ''),
                    interview_data.get('duration', 0),
                    interview_data.get('chunks_total', 0),
                    interview_data.get('chunks_processed', 0),
                    'Sim' if interview_data.get('analysis') else 'Não'
                ]
            }
            
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='Resumo', index=False)
            
            # Transcript sheet
            transcript_lines = interview_data.get('transcript', '').split('\n')
            transcript_data = {
                'Linha': range(1, len(transcript_lines) + 1),
                'Conteúdo': transcript_lines
            }
            pd.DataFrame(transcript_data).to_excel(writer, sheet_name='Transcrição', index=False)
            
            # Analysis sheet (if available)
            if interview_data.get('analysis'):
                analysis_lines = interview_data['analysis'].split('\n')
                analysis_data = {
                    'Linha': range(1, len(analysis_lines) + 1),
                    'Conteúdo': analysis_lines
                }
                pd.DataFrame(analysis_data).to_excel(writer, sheet_name='Análise', index=False)
            
            # Metrics sheet (if available)
            if interview_data.get('metrics'):
                metrics_data = {
                    'Métrica': list(interview_data['metrics'].keys()),
                    'Valor': list(interview_data['metrics'].values())
                }
                pd.DataFrame(metrics_data).to_excel(writer, sheet_name='Métricas', index=False)
        
        return file_path
    
    def _export_batch_as_zip(self, interviews: List[Dict], individual_format: str) -> str:
        """Export multiple interviews as ZIP file"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f"entrevistas_batch_{timestamp}.zip"
        zip_path = os.path.join(self.temp_dir, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for interview in interviews:
                try:
                    # Export individual interview
                    file_path = self.export_interview(interview, individual_format, True, True)
                    
                    # Add to ZIP with organized name
                    interview_id = interview.get('id', 'unknown')
                    archive_name = f"entrevista_{interview_id}.{individual_format}"
                    zipf.write(file_path, archive_name)
                    
                    # Clean up individual file
                    os.remove(file_path)
                    
                except Exception as e:
                    SecureLogger.safe_log_error(logger, "Failed to add interview to batch", e, {
                        'interview_id': interview.get('id'),
                        'format': individual_format
                    })
        
        return zip_path
    
    def _export_batch_as_xlsx(self, interviews: List[Dict]) -> str:
        """Export multiple interviews as single Excel file with multiple sheets"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevistas_batch_{timestamp}.xlsx"
        file_path = os.path.join(self.temp_dir, filename)
        
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # Summary sheet with all interviews
            summary_data = []
            for interview in interviews:
                summary_data.append({
                    'ID': interview.get('id', ''),
                    'Data/Hora': interview.get('created_at', ''),
                    'Status': interview.get('status', ''),
                    'Duração': interview.get('duration', 0),
                    'Chunks Total': interview.get('chunks_total', 0),
                    'Chunks Processados': interview.get('chunks_processed', 0),
                    'Possui Análise': 'Sim' if interview.get('analysis') else 'Não',
                    'Telefone': SecureLogger.mask_phone_number(interview.get('phone_number', ''))
                })
            
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='Resumo Geral', index=False)
            
            # Individual sheets for each interview (limited to first 10 to avoid Excel limits)
            for i, interview in enumerate(interviews[:10]):
                sheet_name = f"Entrevista_{i+1}"
                interview_data = {
                    'Campo': ['ID', 'Status', 'Transcrição', 'Análise'],
                    'Valor': [
                        interview.get('id', ''),
                        interview.get('status', ''),
                        interview.get('transcript', '')[:1000] + '...' if len(interview.get('transcript', '')) > 1000 else interview.get('transcript', ''),
                        interview.get('analysis', '')[:1000] + '...' if len(interview.get('analysis', '')) > 1000 else interview.get('analysis', '')
                    ]
                }
                pd.DataFrame(interview_data).to_excel(writer, sheet_name=sheet_name, index=False)
        
        return file_path
    
    def _export_batch_as_csv(self, interviews: List[Dict]) -> str:
        """Export multiple interviews as single CSV file"""
        data = []
        for interview in interviews:
            row = {
                'interview_id': interview.get('id', ''),
                'created_at': interview.get('created_at', ''),
                'status': interview.get('status', ''),
                'duration_minutes': interview.get('duration', 0),
                'chunks_total': interview.get('chunks_total', 0),
                'chunks_processed': interview.get('chunks_processed', 0),
                'transcript_length': len(interview.get('transcript', '')),
                'has_analysis': bool(interview.get('analysis')),
                'phone_number': SecureLogger.mask_phone_number(interview.get('phone_number', ''))
            }
            
            # Add metrics if available
            if interview.get('metrics'):
                for key, value in interview['metrics'].items():
                    row[f'metric_{key}'] = value
            
            data.append(row)
        
        df = pd.DataFrame(data)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"entrevistas_batch_{timestamp}.csv"
        file_path = os.path.join(self.temp_dir, filename)
        
        df.to_csv(file_path, index=False, encoding='utf-8')
        
        return file_path
    
    def _add_metadata_section(self, doc: Document, interview_data: Dict):
        """Add metadata section to Word document"""
        doc.add_heading('Informações Gerais', level=1)
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Table data
        table_data = [
            ('ID da Entrevista', interview_data.get('id', 'N/A')),
            ('Data/Hora', interview_data.get('created_at', 'N/A')),
            ('Status', interview_data.get('status', 'N/A')),
            ('Duração', f"{interview_data.get('duration', 0)} minutos"),
            ('Total de Chunks', str(interview_data.get('chunks_total', 0))),
            ('Chunks Processados', str(interview_data.get('chunks_processed', 0)))
        ]
        
        for i, (label, value) in enumerate(table_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph("")
    
    def _add_transcript_section(self, doc: Document, interview_data: Dict):
        """Add transcript section to Word document"""
        doc.add_heading('Transcrição Completa', level=1)
        
        transcript = interview_data.get('transcript', 'Transcrição não disponível')
        
        # Add transcript with formatting
        for line in transcript.split('\n'):
            if line.strip():
                if any(speaker in line for speaker in ['ENTREVISTADOR:', 'CANDIDATO:', 'LOCUTOR:']):
                    para = doc.add_paragraph()
                    speaker_end = line.find(':')
                    if speaker_end > 0:
                        run = para.add_run(line[:speaker_end + 1])
                        run.bold = True
                        para.add_run(line[speaker_end + 1:])
                    else:
                        para.add_run(line)
                else:
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph("")
        
        doc.add_paragraph("")
    
    def _add_analysis_section(self, doc: Document, interview_data: Dict):
        """Add analysis section to Word document"""
        doc.add_heading('Análise da Entrevista', level=1)
        
        analysis = interview_data.get('analysis', '')
        
        # Parse and format analysis
        sections = analysis.split('**')
        current_text = ""
        
        for i, section in enumerate(sections):
            if i % 2 == 1:  # Heading
                if current_text.strip():
                    doc.add_paragraph(current_text.strip())
                    current_text = ""
                doc.add_heading(section.strip(), level=2)
            else:
                current_text += section
        
        if current_text.strip():
            doc.add_paragraph(current_text.strip())
        
        doc.add_paragraph("")
    
    def _add_metrics_section(self, doc: Document, interview_data: Dict):
        """Add performance metrics section to Word document"""
        doc.add_heading('Métricas de Performance', level=1)
        
        metrics = interview_data.get('metrics', {})
        
        if metrics:
            table = doc.add_table(rows=len(metrics) + 1, cols=2)
            table.style = 'Table Grid'
            
            # Header
            table.cell(0, 0).text = 'Métrica'
            table.cell(0, 1).text = 'Valor'
            table.cell(0, 0).paragraphs[0].runs[0].bold = True
            table.cell(0, 1).paragraphs[0].runs[0].bold = True
            
            # Data
            for i, (key, value) in enumerate(metrics.items(), 1):
                table.cell(i, 0).text = key
                table.cell(i, 1).text = str(value)
        else:
            doc.add_paragraph("Nenhuma métrica de performance disponível.")
        
        doc.add_paragraph("")

# Global export manager instance
export_manager = ExportManager()